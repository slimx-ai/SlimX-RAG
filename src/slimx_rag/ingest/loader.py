from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable, List, Optional

from langchain_core.documents import Document

from slimx_rag.settings import IndexingPipelineSettings
from slimx_rag.core.hashing import content_hash, fallback_doc_id, path_id

logger = logging.getLogger(__name__)


def iter_subdirs(kb_dir: Path) -> Iterable[Path]:
    """Yield subdirectories of kb_dir representing document types."""
    if not kb_dir.exists():
        raise FileNotFoundError(f"Knowledge base directory not found: {kb_dir}")
    if not kb_dir.is_dir():
        raise NotADirectoryError(f"Knowledge base path is not a directory: {kb_dir}")

    for p in kb_dir.iterdir():
        if p.is_dir():
            yield p


def _doc_type_from_relpath(relpath: Path, depth: int) -> str:
    """Derive doc_type from relative path parts up to the given depth."""
    parts = relpath.parts[: max(1, depth)]
    return "/".join(parts) if parts else ""


def _read_text_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt", ".text"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception as e:  # pragma: no cover - optional dependency
            raise RuntimeError("PDF ingestion requires optional dependency 'pypdf'.") from e
        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".docx":
        try:
            import docx  # type: ignore
        except Exception as e:  # pragma: no cover - optional dependency
            raise RuntimeError("DOCX ingestion requires optional dependency 'python-docx'.") from e
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    return path.read_text(encoding="utf-8")


def _fallback_doc_id(source: str, content_hash_value: str) -> str:
    # identity for docs with no stable kb_relpath: stable for same source+content
    return fallback_doc_id(source, content_hash_value)


def fetch_documents(settings: IndexingPipelineSettings) -> List[Document]:
    """
    Load documents from knowledge-base (subfolders) and attach stable metadata.

    Baseline metadata (always set):
      - doc_id          (stable identity; relpath-based when possible)
      - content_hash    (version fingerprint; changes when content changes)
      - content_len     (chars)
      - doc_type        (folder-derived when possible, else "unknown")
      - file_ext        (when discoverable)
      - kb_relpath      (when discoverable)

    Notes:
      - doc_id is *identity* (path-based when possible)
      - content_hash is *version* (content-based)
      - content_len counts the exact loaded text, including final newline characters
    """

    settings.ingest.validate()
    if not settings.kb_dir.exists():
        raise FileNotFoundError(f"kb_dir not found: {settings.kb_dir}")
    if not settings.kb_dir.is_dir():
        raise NotADirectoryError(f"kb_dir is not a directory: {settings.kb_dir}")


    kb_dir = settings.kb_dir
    glob = settings.ingest.glob
    documents: List[Document] = []
    logger.info(f"Loading documents from knowledge base at: {kb_dir}")

    paths = sorted(p for p in kb_dir.glob(glob) if p.is_file())
    for source_path in paths:
        relpath = source_path.relative_to(kb_dir)
        text = _read_text_path(source_path)
        metadata = {
            "source": str(source_path),
            "title": source_path.stem,
            "content_hash": content_hash(text),
            "content_len": len(text),
            "kb_relpath": relpath.as_posix(),
            "file_ext": relpath.suffix.lower().lstrip("."),
            "doc_id": path_id(relpath.as_posix()),
            "doc_type": (
                _doc_type_from_relpath(relpath, depth=settings.ingest.doc_type_depth)
                if settings.ingest.doc_type_mode == "subdir"
                else "unknown"
            ),
        }
        if not metadata["doc_id"]:
            metadata["doc_id"] = _fallback_doc_id(str(source_path), metadata["content_hash"])
        documents.append(Document(page_content=text, metadata=metadata))

    logger.info(f"Fetched {len(documents)} documents from knowledge base.")
    return documents
